from collections.abc import Generator
from typing import Any
import os
import re
import json

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from dify_plugin.entities.model.message import UserPromptMessage

from docx import Document
from docx.oxml.ns import qn

from tools.utils import clean_paths, save_upload_to_temp, strip_model_thoughts, invoke_llm, dual_messages


class FileRevisionTool(Tool):
    @staticmethod
    def _severity_rank(sev: str) -> int:
        sval = str(sev or "").strip().lower()
        if sval == "high":
            return 3
        if sval == "medium":
            return 2
        if sval == "low":
            return 1
        return 0

    @staticmethod
    def _extract_comment_anchors(doc: Document) -> dict[int, dict[str, Any]]:
        start_tag = qn("w:commentRangeStart")
        end_tag = qn("w:commentRangeEnd")
        run_tag = qn("w:r")
        text_tag = qn("w:t")
        id_attr = qn("w:id")

        by_id: dict[int, dict[str, Any]] = {}
        for pid, para in enumerate(doc.paragraphs):
            active_ids: set[int] = set()
            for child in list(para._p.iterchildren()):
                if child.tag == start_tag:
                    raw_id = child.get(id_attr)
                    try:
                        cid = int(raw_id) if raw_id is not None else None
                    except Exception:
                        cid = None
                    if cid is not None:
                        active_ids.add(cid)
                        by_id.setdefault(cid, {"chunks": [], "para_ids": set()})
                        by_id[cid]["para_ids"].add(pid)
                    continue

                if child.tag == end_tag:
                    raw_id = child.get(id_attr)
                    try:
                        cid = int(raw_id) if raw_id is not None else None
                    except Exception:
                        cid = None
                    if cid is not None and cid in active_ids:
                        active_ids.remove(cid)
                    continue

                if child.tag == run_tag and active_ids:
                    run_text = "".join([(t.text or "") for t in child.iter() if t.tag == text_tag])
                    if run_text:
                        for cid in list(active_ids):
                            by_id.setdefault(cid, {"chunks": [], "para_ids": set()})
                            by_id[cid]["chunks"].append(run_text)
                            by_id[cid]["para_ids"].add(pid)

        out: dict[int, dict[str, Any]] = {}
        for cid, info in by_id.items():
            out[cid] = {
                "anchor_text": "".join(info.get("chunks", [])).strip(),
                "para_ids": sorted([int(x) for x in info.get("para_ids", set())]),
            }
        return out

    @staticmethod
    def _parse_comment_payload(comment_text: str) -> dict[str, str]:
        text = str(comment_text or "").strip()
        lines = text.splitlines()
        first = lines[0].strip() if lines else ""

        code = ""
        severity = "medium"
        comment_body = ""

        m = re.match(r"^\[([^\]]+)\]\[([^\]]+)\]\s*(.*)$", first)
        if m:
            code = m.group(1).strip()
            severity = m.group(2).strip().lower()
            comment_body = m.group(3).strip()

        original = ""
        revised = ""
        mo = re.search(r"【(?:原文|Original)】：(.*?)(?:\n【(?:修改后|After modification)】：(.*))?$", text, flags=re.DOTALL)
        if mo:
            original = (mo.group(1) or "").strip()
            revised = (mo.group(2) or "").strip()

        return {
            "code": code,
            "severity": severity if severity in ["high", "medium", "low"] else "medium",
            "comment_body": comment_body,
            "original": original,
            "revised": revised,
            "raw": text,
        }

    @staticmethod
    def _remove_comment_elements(doc: Document, remove_ids: set[int]) -> None:
        if not remove_ids:
            return

        comments = list(doc.comments)
        for c in comments:
            if c.comment_id in remove_ids:
                elm = c._comment_elm
                parent = elm.getparent()
                if parent is not None:
                    parent.remove(elm)

        comment_tags = {
            qn("w:commentRangeStart"),
            qn("w:commentRangeEnd"),
            qn("w:commentReference"),
        }
        id_attr = qn("w:id")

        for el in list(doc._element.iter()):
            if el.tag in comment_tags:
                raw_id = el.get(id_attr)
                if raw_id is None:
                    continue
                try:
                    cid = int(raw_id)
                except Exception:
                    continue
                if cid in remove_ids:
                    parent = el.getparent()
                    if parent is not None:
                        parent.remove(el)

    @staticmethod
    def _pick_target_paragraph(doc: Document, anchor_text: str, para_ids: list[int] | None = None) -> Any:
        atxt = str(anchor_text or "").strip()
        if para_ids:
            for pid in para_ids:
                if 0 <= int(pid) < len(doc.paragraphs):
                    p = doc.paragraphs[int(pid)]
                    if not atxt or atxt in (p.text or ""):
                        return p
        if atxt:
            for p in doc.paragraphs:
                if atxt in (p.text or ""):
                    return p
        for p in doc.paragraphs:
            if (p.text or "").strip():
                return p
        return doc.paragraphs[0] if doc.paragraphs else None

    @staticmethod
    def _pick_target_run(para: Any, anchor_text: str) -> Any:
        if para is None:
            return None
        runs = list(para.runs)
        if not runs:
            r = para.add_run(para.text or "")
            runs = [r]
        atxt = str(anchor_text or "").strip()
        if atxt:
            for r in runs:
                if atxt in (r.text or ""):
                    return r
        for r in runs:
            if (r.text or "").strip():
                return r
        return runs[0] if runs else None

    def _semantic_choose_one(self, llm_model: dict[str, Any], entries: list[dict[str, Any]]) -> int:
        options = []
        for idx, e in enumerate(entries, start=1):
            options.append(
                f"{idx}. [rule_code={e['code']}][severity={e['severity']}] {e['comment_body']} | original={e['original']} | revised={e['revised']}"
            )
        prompt = (
            "You are a legal review merge assistant. Choose the most representative single annotation from options.\n"
            "Return JSON only: {\"keep_index\": <int>}\n"
            "Options:\n" + "\n".join(options)
        )
        try:
            resp = invoke_llm(self, llm_model, [UserPromptMessage(content=prompt)])
            parsed = json.loads(strip_model_thoughts(resp))
            keep_index = int(parsed.get("keep_index", 1))
            if 1 <= keep_index <= len(entries):
                return keep_index - 1
        except Exception:
            pass
        return 0

    def _semantic_merge(self, llm_model: dict[str, Any], entries: list[dict[str, Any]]) -> dict[str, str]:
        options = []
        for idx, e in enumerate(entries, start=1):
            options.append(
                f"{idx}. [rule_code={e['code']}][severity={e['severity']}] {e['comment_body']} | original={e['original']} | revised={e['revised']}"
            )
        prompt = (
            "You are a legal review merge assistant. Merge multi-risk annotations into one improved annotation.\n"
            "Return JSON only:\n"
            "{\"severity\":\"high|medium|low\",\"comment_text\":\"...\",\"revised\":\"...\"}\n"
            "Requirements:\n"
            "1) keep revised concise and directly usable.\n"
            "2) severity should reflect final merged risk level.\n"
            "3) do not output markdown.\n"
            "Inputs:\n" + "\n".join(options)
        )
        try:
            resp = invoke_llm(self, llm_model, [UserPromptMessage(content=prompt)])
            parsed = json.loads(strip_model_thoughts(resp))
            sev = str(parsed.get("severity", "")).strip().lower()
            if sev not in ["high", "medium", "low"]:
                sev = "medium"
            comment_text = str(parsed.get("comment_text", "")).strip()
            revised = str(parsed.get("revised", "")).strip()
            return {
                "severity": sev,
                "comment_body": comment_text,
                "revised": revised,
            }
        except Exception:
            return {
                "severity": "medium",
                "comment_body": entries[0]["comment_body"] if entries else "",
                "revised": entries[0]["revised"] if entries else "",
            }

    def _second_pass_overlap_merge(
        self,
        llm_model: dict[str, Any],
        entries: list[dict[str, Any]],
        doc: Document,
    ) -> tuple[list[dict[str, Any]], int]:
        index_map = {idx: e for idx, e in enumerate(entries)}
        visited: set[int] = set()
        final: list[dict[str, Any]] = []
        merged_groups = 0

        para_to_indices: dict[int, list[int]] = {}
        for idx, e in index_map.items():
            pids = [int(x) for x in e.get("para_ids", []) if isinstance(x, int) or str(x).isdigit()]
            for pid in pids:
                para_to_indices.setdefault(pid, []).append(idx)

        def _span_in_para(para_text: str, src: str) -> tuple[int, int] | None:
            s = str(src or "").strip()
            if not s:
                return None
            start = para_text.find(s)
            if start >= 0:
                return (start, start + len(s))
            return None

        def _connected(i: int, j: int, pid: int) -> bool:
            ei = index_map[i]
            ej = index_map[j]
            para_text = doc.paragraphs[pid].text or ""
            ti = str(ei.get("original") or ei.get("anchor_text") or "").strip()
            tj = str(ej.get("original") or ej.get("anchor_text") or "").strip()
            si = _span_in_para(para_text, ti)
            sj = _span_in_para(para_text, tj)
            if si and sj:
                return not (si[1] <= sj[0] or sj[1] <= si[0])
            if ti and tj and (ti in tj or tj in ti):
                return True
            return False

        adjacency: dict[int, set[int]] = {idx: set() for idx in index_map.keys()}
        for pid, idx_list in para_to_indices.items():
            uniq = sorted(set(idx_list))
            for a in range(len(uniq)):
                for b in range(a + 1, len(uniq)):
                    i = uniq[a]
                    j = uniq[b]
                    if _connected(i, j, pid):
                        adjacency[i].add(j)
                        adjacency[j].add(i)

        for idx in sorted(index_map.keys()):
            if idx in visited:
                continue
            comp: list[int] = []
            stack = [idx]
            visited.add(idx)
            while stack:
                cur = stack.pop()
                comp.append(cur)
                for nxt in adjacency.get(cur, set()):
                    if nxt not in visited:
                        visited.add(nxt)
                        stack.append(nxt)

            if len(comp) == 1:
                final.append(index_map[comp[0]])
                continue

            merged_groups += 1
            cluster = [index_map[c] for c in comp]
            merged = self._semantic_merge(llm_model, cluster)
            raw_codes: list[str] = []
            for c in cluster:
                one_code = str(c.get("code") or "").strip()
                if not one_code:
                    continue
                for part in one_code.replace(",", "、").split("、"):
                    p = part.strip()
                    if p:
                        raw_codes.append(p)
            code_join = "、".join(dict.fromkeys(raw_codes)) or "NO_RULE"
            highest = "low"
            para_ids_union: set[int] = set()
            for c in cluster:
                if self._severity_rank(c.get("severity")) > self._severity_rank(highest):
                    highest = str(c.get("severity") or "low")
                for pid in c.get("para_ids", []):
                    if isinstance(pid, int) or str(pid).isdigit():
                        para_ids_union.add(int(pid))
            source_candidates = [str(c.get("original") or c.get("anchor_text") or "").strip() for c in cluster]
            source_candidates = [x for x in source_candidates if x]
            source_text = sorted(source_candidates, key=lambda x: len(x), reverse=True)[0] if source_candidates else ""

            final.append(
                {
                    "comment_id": None,
                    "code": code_join,
                    "severity": merged.get("severity") or highest,
                    "comment_body": merged.get("comment_body") or cluster[0].get("comment_body", ""),
                    "original": source_text,
                    "revised": merged.get("revised") or cluster[0].get("revised", ""),
                    "anchor_text": source_text,
                    "para_ids": sorted(para_ids_union),
                    "raw": "",
                }
            )

        return final, merged_groups

    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        llm_model = tool_parameters.get("model_config")
        file_obj = tool_parameters.get("upload_file")
        merge_strategy = str(tool_parameters.get("merge_strategy") or "keep_highest_risk").strip().lower()
        apply_to_original = str(tool_parameters.get("apply_to_original") or "no").strip().lower()
        output_file_name = tool_parameters.get("output_file_name")

        if not isinstance(llm_model, dict):
            for m in dual_messages(self, "Error: model_config invalid.", {"error": "model_config invalid"}):
                yield m
            return
        if not file_obj:
            for m in dual_messages(self, "Error: No file uploaded.", {"error": "No file uploaded"}):
                yield m
            return

        if merge_strategy not in ["keep_highest_risk", "keep_semantic", "merge_semantic"]:
            merge_strategy = "keep_highest_risk"
        apply_changes = apply_to_original in ["yes", "true", "1"]

        temp_path = None
        try:
            temp_path, original_name, ext = save_upload_to_temp(file_obj)
            if ext != ".docx":
                for m in dual_messages(self, "Error: Only .docx is supported.", {"error": "Only .docx is supported"}):
                    yield m
                return

            if not output_file_name or not str(output_file_name).strip():
                base, _ = os.path.splitext(original_name)
                output_file_name = f"revised_{base}"

            doc = Document(temp_path)
            anchor_map = self._extract_comment_anchors(doc)

            managed_comments: list[dict[str, Any]] = []
            for c in list(doc.comments):
                parsed = self._parse_comment_payload(c.text)
                if parsed["code"]:
                    one = {
                        "comment_id": c.comment_id,
                        "anchor_text": anchor_map.get(c.comment_id, {}).get("anchor_text", ""),
                        "para_ids": anchor_map.get(c.comment_id, {}).get("para_ids", []),
                        **parsed,
                    }
                    managed_comments.append(one)

            if not managed_comments:
                payload = {
                    "status": "ok",
                    "output_file": f"{output_file_name}{ext}",
                    "merge_strategy": merge_strategy,
                    "apply_to_original": apply_changes,
                    "managed_comment_count": 0,
                    "merged_group_count": 0,
                    "final_comment_count": 0,
                    "modified_count": 0,
                }
                output_path = os.path.join(os.path.dirname(temp_path), f"{output_file_name}{ext}")
                doc.save(output_path)
                with open(output_path, "rb") as f:
                    data = f.read()
                for m in dual_messages(self, json.dumps(payload, ensure_ascii=False), payload):
                    yield m
                yield self.create_blob_message(
                    blob=data,
                    meta={
                        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "save_as": f"{output_file_name}{ext}",
                        "filename": f"{output_file_name}{ext}",
                    },
                )
                return

            groups: dict[str, list[dict[str, Any]]] = {}
            for item in managed_comments:
                key = (item.get("original") or "").strip()
                if not key:
                    key = f"__single_{item['comment_id']}"
                groups.setdefault(key, []).append(item)

            final_entries: list[dict[str, Any]] = []
            merged_group_count = 0

            for gkey, items in groups.items():
                if len(items) == 1 or gkey.startswith("__single_"):
                    final_entries.append(items[0])
                    continue

                merged_group_count += 1
                if merge_strategy == "keep_highest_risk":
                    ranks = [self._severity_rank(x.get("severity")) for x in items]
                    top_rank = max(ranks) if ranks else 0
                    top_items = [x for x in items if self._severity_rank(x.get("severity")) == top_rank]
                    if len(top_items) == 1:
                        final_entries.append(top_items[0])
                    else:
                        keep_idx = self._semantic_choose_one(llm_model, top_items)
                        final_entries.append(top_items[keep_idx])
                elif merge_strategy == "keep_semantic":
                    keep_idx = self._semantic_choose_one(llm_model, items)
                    final_entries.append(items[keep_idx])
                else:
                    merged = self._semantic_merge(llm_model, items)
                    code_join = "、".join(dict.fromkeys([x.get("code", "").strip() for x in items if x.get("code")]))
                    highest = "low"
                    for x in items:
                        if self._severity_rank(x.get("severity")) > self._severity_rank(highest):
                            highest = x.get("severity", "low")
                    final_entries.append(
                        {
                            "comment_id": None,
                            "code": code_join or "NO_RULE",
                            "severity": merged.get("severity") or highest,
                            "comment_body": merged.get("comment_body") or items[0].get("comment_body", ""),
                            "original": items[0].get("original", ""),
                            "revised": merged.get("revised") or items[0].get("revised", ""),
                            "anchor_text": items[0].get("anchor_text", ""),
                            "para_ids": items[0].get("para_ids", []),
                            "raw": "",
                        }
                    )

            final_entries, overlap_merged_group_count = self._second_pass_overlap_merge(llm_model, final_entries, doc)

            remove_ids = {int(x["comment_id"]) for x in managed_comments if x.get("comment_id") is not None}
            self._remove_comment_elements(doc, remove_ids)

            modified_count = 0
            if apply_changes:
                final_entries = sorted(
                    final_entries,
                    key=lambda x: len(str(x.get("original") or x.get("anchor_text") or "")),
                    reverse=True,
                )
                for item in final_entries:
                    original = str(item.get("original") or item.get("anchor_text") or "").strip()
                    revised = str(item.get("revised") or "").strip()
                    if not original or not revised or original == revised:
                        continue
                    para_ids = [int(x) for x in item.get("para_ids", []) if isinstance(x, int) or str(x).isdigit()]
                    para = self._pick_target_paragraph(doc, original, para_ids)
                    if para is None:
                        continue
                    changed = False
                    tr = self._pick_target_run(para, original)
                    if tr is not None and original in (tr.text or ""):
                        tr.text = (tr.text or "").replace(original, revised, 1)
                        changed = True
                    elif original in (para.text or ""):
                        para.text = (para.text or "").replace(original, revised, 1)
                        changed = True
                    else:
                        for p in doc.paragraphs:
                            if original in (p.text or ""):
                                p.text = (p.text or "").replace(original, revised, 1)
                                changed = True
                                break
                    if changed:
                        modified_count += 1

            for item in final_entries:
                code = str(item.get("code") or "NO_RULE").strip()
                severity = str(item.get("severity") or "medium").strip().lower()
                if severity not in ["high", "medium", "low"]:
                    severity = "medium"
                comment_body = str(item.get("comment_body") or "").strip() or "Risk detected."
                original = str(item.get("original") or item.get("anchor_text") or "").strip()
                revised = str(item.get("revised") or "").strip()

                final_comment = f"[{code}][{severity}] {comment_body}"
                detail = f"【原文】：{original}\n【修改后】：{revised}"
                full_comment = f"{final_comment}\n{detail}"

                anchor = revised if (apply_changes and revised) else original
                para_ids = [int(x) for x in item.get("para_ids", []) if isinstance(x, int) or str(x).isdigit()]
                para = self._pick_target_paragraph(doc, anchor, para_ids)
                if para is None:
                    continue
                tr = self._pick_target_run(para, anchor)
                if tr is None:
                    continue
                doc.add_comment([tr], full_comment, author="DocReview", initials="DR")

            output_path = os.path.join(os.path.dirname(temp_path), f"{output_file_name}{ext}")
            doc.save(output_path)
            with open(output_path, "rb") as f:
                data = f.read()

            summary_payload = {
                "status": "ok",
                "output_file": f"{output_file_name}{ext}",
                "merge_strategy": merge_strategy,
                "apply_to_original": apply_changes,
                "managed_comment_count": len(managed_comments),
                "merged_group_count": merged_group_count,
                "overlap_merged_group_count": overlap_merged_group_count,
                "final_comment_count": len(final_entries),
                "modified_count": modified_count,
            }
            for m in dual_messages(self, json.dumps(summary_payload, ensure_ascii=False), summary_payload):
                yield m
            yield self.create_blob_message(
                blob=data,
                meta={
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "save_as": f"{output_file_name}{ext}",
                    "filename": f"{output_file_name}{ext}",
                },
            )

        finally:
            clean_paths([temp_path] if temp_path else [])
