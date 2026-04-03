from collections.abc import Generator
from typing import Any
import os
import json
import hashlib

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from dify_plugin.entities.model.message import UserPromptMessage

from tools.utils import clean_paths, save_upload_to_temp, strip_model_thoughts, invoke_llm, dual_messages, safe_json_load, detect_text_language
from docx import Document


class DocAnnotatorTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        llm_model = tool_parameters.get("model_config")
        file_obj = tool_parameters.get("upload_file")
        audit_report = tool_parameters.get("audit_report") or ""
        output_file_name = tool_parameters.get("output_file_name")
        annotation_style = tool_parameters.get("annotation_style") or "comment"
        apply_to_original = str(tool_parameters.get("apply_to_original") or "no").strip().lower()
        apply_changes = apply_to_original in ["yes", "true", "1"]

        if not isinstance(llm_model, dict):
            for m in dual_messages(self, "Error: model_config invalid.", {"error": "model_config invalid"}):
                yield m
            return
        if not file_obj:
            for m in dual_messages(self, "Error: No file uploaded.", {"error": "No file uploaded"}):
                yield m
            return

        temp_path = None
        try:
            temp_path, original_name, ext = save_upload_to_temp(file_obj)
            if ext != ".docx":
                for m in dual_messages(self, "Error: Only .docx is supported.", {"error": "Only .docx is supported"}):
                    yield m
                return

            if not output_file_name or not str(output_file_name).strip():
                base, _ = os.path.splitext(original_name)
                output_file_name = f"reviewed_{base}"

            if isinstance(audit_report, dict):
                report_payload = audit_report
            else:
                report_payload = safe_json_load(audit_report, {})
            if isinstance(report_payload, str):
                report_payload = safe_json_load(report_payload, {})
            output_language = ""
            lang_from_report = ""
            if isinstance(report_payload, dict):
                lang_from_report = str(
                    report_payload.get("output_language")
                    or report_payload.get("language")
                    or report_payload.get("summary", {}).get("output_language", "")
                ).strip().lower()
            if lang_from_report in ["zh", "en", "ja", "ko", "es", "fr", "de", "pt", "ru", "ar"]:
                output_language = lang_from_report
            else:
                doc = Document(temp_path)
                probe = "\n".join([p.text for p in doc.paragraphs[:20]])
                output_language = detect_text_language(probe)
            if output_language not in ["zh", "en", "ja", "ko", "es", "fr", "de", "pt", "ru", "ar"]:
                output_language = "en"

            def _collect_risks(obj: Any) -> list[dict[str, Any]]:
                out: list[dict[str, Any]] = []
                if isinstance(obj, str):
                    parsed = safe_json_load(obj, None)
                    if parsed is not None:
                        out.extend(_collect_risks(parsed))
                    return out
                if isinstance(obj, list):
                    for one in obj:
                        out.extend(_collect_risks(one))
                    return out
                if isinstance(obj, dict):
                    if isinstance(obj.get("risks"), list):
                        out.extend([x for x in obj.get("risks", []) if isinstance(x, dict)])
                    if isinstance(obj.get("audit_results"), list):
                        out.extend([x for x in obj.get("audit_results", []) if isinstance(x, dict)])
                    if isinstance(obj.get("json"), list):
                        for j in obj.get("json", []):
                            out.extend(_collect_risks(j))
                    if isinstance(obj.get("text"), str):
                        out.extend(_collect_risks(obj.get("text")))
                return out

            risks = _collect_risks(report_payload)
            if not risks:
                for m in dual_messages(self, "Error: audit_report has no risks/audit_results list.", {"error": "audit_report has no risks/audit_results list"}):
                    yield m
                    return

            doc = Document(temp_path)
            para_map = {idx: p for idx, p in enumerate(doc.paragraphs)}
            para_hash_map = {
                idx: hashlib.sha1((p.text or "").strip().encode("utf-8")).hexdigest()[:12]
                for idx, p in para_map.items()
            }
            hash_to_pids: dict[str, list[int]] = {}
            for idx, ph in para_hash_map.items():
                hash_to_pids.setdefault(ph, []).append(idx)
            annotation_count = 0
            skipped_count = 0
            mislocated_count = 0
            modified_count = 0
            located_by_quote = 0
            located_by_ref = 0
            located_by_chunk = 0
            located_by_hash = 0

            placed_count_by_pid: dict[int, int] = {}

            for item in risks:
                if not isinstance(item, dict):
                    continue

                code = str(item.get("matched_rule_code", "")).strip() or "NO_RULE"
                severity = str(item.get("severity", "")).strip().lower() or "medium"
                if severity not in ["high", "medium", "low"]:
                    severity = "medium"

                reason = str(item.get("reason", "")).strip()
                suggestion = str(item.get("suggestion", "")).strip()
                quote = str(item.get("quote", "")).strip()

                comment_prompt = f"""
You are a legal annotation assistant.
Language: {output_language}

Risk info:
- rule_code: {code}
- severity: {severity}
- quote: {quote}
- reason: {reason}
- suggestion: {suggestion}

Task:
Generate ONE concise annotation sentence (no JSON) in {output_language}.
"""

                try:
                    result = invoke_llm(self, llm_model, [UserPromptMessage(content=comment_prompt)])
                    comment_text = strip_model_thoughts(result).strip()
                except Exception as e:
                    for m in dual_messages(self, f"LLM Error: {str(e)}", {"error": f"LLM Error: {str(e)}"}):
                        yield m
                    return

                if not comment_text:
                    comment_text = reason or suggestion or "Risk detected."

                rewrite_prompt = f"""
You are a contract editing assistant.
Language: {output_language}

Original text:
{quote}

Revision instruction:
{suggestion}

Task:
Return one revised replacement text for the original text, concise and directly usable.
Output plain text only.
"""
                revised_text = ""
                if quote and suggestion:
                    try:
                        revised_resp = invoke_llm(self, llm_model, [UserPromptMessage(content=rewrite_prompt)])
                        revised_text = strip_model_thoughts(revised_resp).strip()
                    except Exception:
                        revised_text = ""
                if not revised_text:
                    revised_text = suggestion or quote

                para = None
                pid = None
                refs = item.get("element_refs", [])
                candidate_pids: list[int] = []
                candidate_hashes: list[str] = []

                emeta = item.get("element_meta", [])
                if isinstance(emeta, list):
                    for meta_item in emeta:
                        if isinstance(meta_item, dict):
                            ph = str(meta_item.get("para_hash", "")).strip()
                            if ph:
                                candidate_hashes.append(ph)

                if candidate_hashes:
                    for ph in candidate_hashes:
                        for cand in hash_to_pids.get(ph, []):
                            if cand not in candidate_pids:
                                candidate_pids.append(cand)

                if isinstance(refs, list) and refs:
                    for one_ref in refs:
                        sref = str(one_ref)
                        if sref.startswith("p:"):
                            try:
                                cand = int(sref.split(":", 1)[1])
                            except Exception:
                                continue
                            if cand in para_map:
                                candidate_pids.append(cand)
                            used = placed_count_by_pid.get(cand, 0)
                            if used == 0 and cand in para_map:
                                pid = cand
                                break

                    if pid is None:
                        for one_ref in refs:
                            sref = str(one_ref)
                            if sref.startswith("p:"):
                                try:
                                    cand = int(sref.split(":", 1)[1])
                                except Exception:
                                    continue
                                if cand in para_map:
                                    pid = cand
                                    break

                if pid is None:
                    chunk_id_val = item.get("chunk_id")
                    if chunk_id_val is not None:
                        try:
                            cand = int(chunk_id_val)
                            if cand in para_map:
                                pid = cand
                        except Exception:
                            pid = None

                pid_para_text = (para_map[pid].text or "") if (pid is not None and pid in para_map) else ""
                if quote and pid is not None and quote not in pid_para_text:
                    found_pid = None
                    for cand_idx, cand_para in para_map.items():
                        if quote in (cand_para.text or ""):
                            found_pid = cand_idx
                            break
                    if found_pid is not None:
                        pid = found_pid

                if pid is not None:
                    para = para_map.get(pid)

                if not para:
                    continue

                # Prefer paragraph that really contains quote text.
                if quote and candidate_pids:
                    for cand in candidate_pids:
                        ptxt = para_map[cand].text or ""
                        if quote in ptxt:
                            para = para_map[cand]
                            pid = cand
                            break

                if para is None and candidate_hashes:
                    for ph in candidate_hashes:
                        for cand in hash_to_pids.get(ph, []):
                            para = para_map.get(cand)
                            pid = cand
                            if para is not None:
                                break
                        if para is not None:
                            break

                if para is None:
                    skipped_count += 1
                    continue

                if quote and pid is not None and quote in (para.text or ""):
                    located_by_quote += 1
                elif pid is not None and candidate_hashes:
                    located_by_hash += 1
                elif pid is not None and candidate_pids:
                    located_by_ref += 1
                elif pid is not None:
                    located_by_chunk += 1

                # Hard validation: if original quote not in anchor paragraph, skip writing comment.
                if quote and quote not in (para.text or ""):
                    mislocated_count += 1
                    skipped_count += 1
                    continue

                final_comment = f"[{code}][{severity}] {comment_text}"
                runs = list(para.runs)
                if not runs:
                    run = para.add_run(para.text)
                    runs = [run]

                target_run = None
                if quote:
                    for r in runs:
                        rtxt = r.text or ""
                        if rtxt and quote in rtxt:
                            target_run = r
                            break
                if target_run is None:
                    for r in runs:
                        if (r.text or "").strip():
                            target_run = r
                            break
                if target_run is None:
                    target_run = runs[0]

                # Optionally modify source text.
                if apply_changes and quote and revised_text:
                    changed = False
                    if target_run is not None and quote in (target_run.text or ""):
                        target_run.text = (target_run.text or "").replace(quote, revised_text, 1)
                        changed = True
                    elif quote in (para.text or ""):
                        para.text = (para.text or "").replace(quote, revised_text, 1)
                        runs2 = list(para.runs)
                        if runs2:
                            target_run = runs2[0]
                        changed = True
                    if changed:
                        modified_count += 1

                if output_language == "zh":
                    original_label = "原文"
                    after_label = "修改后"
                else:
                    original_label = "Original"
                    after_label = "After modification"

                original_for_comment = quote or (para.text or "")[:120]
                after_for_comment = revised_text or suggestion or original_for_comment
                detail_block = f"【{original_label}】：{original_for_comment}\n【{after_label}】：{after_for_comment}"

                full_comment = f"{final_comment}\n{detail_block}"
                doc.add_comment([target_run], full_comment, author="DocReview", initials="DR")
                if pid is not None:
                    placed_count_by_pid[pid] = placed_count_by_pid.get(pid, 0) + 1
                annotation_count += 1

            output_path = os.path.join(os.path.dirname(temp_path), f"{output_file_name}{ext}")
            doc.save(output_path)
            with open(output_path, "rb") as f:
                data = f.read()

            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            file_name = f"{output_file_name}{ext}"
            summary_payload = {
                "status": "ok",
                "output_file": file_name,
                "annotation_count": annotation_count,
                "modified_count": modified_count,
                "apply_to_original": apply_changes,
                "output_language": output_language,
                "located_by_quote": located_by_quote,
                "located_by_hash": located_by_hash,
                "located_by_ref": located_by_ref,
                "located_by_chunk": located_by_chunk,
                "skipped_count": skipped_count,
                "mislocated_count": mislocated_count,
            }
            for m in dual_messages(self, json.dumps(summary_payload, ensure_ascii=False), summary_payload):
                yield m
            yield self.create_blob_message(blob=data, meta={"mime_type": mime_type, "save_as": file_name, "filename": file_name})

        finally:
            clean_paths([temp_path] if temp_path else [])
